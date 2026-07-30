import json
import logging
from sqlalchemy.orm import Session
import google.generativeai as genai
from app.config.config import get_settings
from app.models.profile import Profile
from app.models.resume import Resume
from app.models.parsed_resume import ParsedResume
from app.models.analysis import Analysis
from app.models.resume_draft import ResumeDraft
from app.schemas.resume_draft import ResumeDraftCreate, ResumeDraftUpdate

# Note: PDF/DOCX generation will be added after libraries install

logger = logging.getLogger(__name__)
settings = get_settings()

if settings.GEMINI_API_KEY:
    genai.configure(api_key=settings.GEMINI_API_KEY)

class ResumeBuilderService:
    def __init__(self):
        self.model = genai.GenerativeModel('gemini-flash-latest')

    def _clean_json(self, text: str) -> dict:
        if text.startswith("```json"):
            text = text.replace("```json", "", 1)
        if text.startswith("```"):
            text = text.replace("```", "", 1)
        if text.endswith("```"):
            text = text[::-1].replace("```", "", 1)[::-1]
        return json.loads(text.strip())

    def get_initial_resume_data(self, db: Session, user_id: str):
        profile = db.query(Profile).filter(Profile.user_id == user_id).first()
        resume = db.query(Resume).filter(Resume.user_id == user_id).order_by(Resume.uploaded_at.desc()).first()
        
        parsed = None
        analysis = None
        
        if resume:
            parsed = db.query(ParsedResume).filter(ParsedResume.resume_id == resume.id).first()
            analysis = db.query(Analysis).filter(Analysis.resume_id == resume.id).first()
        
        # Build initial content structure
        content = {
            "personal_info": {
                "name": profile.full_name if profile else "",
                "title": profile.bio if profile else "",
                "email": getattr(profile, 'user', None).email if profile and getattr(profile, 'user', None) else "",
                "phone": profile.phone if profile else "",
                "location": f"{profile.current_city}, {profile.country}" if profile and profile.current_city else "",
                "github": profile.github_url if profile else "",
                "linkedin": profile.linkedin_url if profile else "",
                "portfolio": profile.portfolio_url if profile else "",
            },
            "summary": profile.bio if profile else "",
            "experience": [],
            "education": [],
            "skills": profile.skills.split(",") if profile and profile.skills else [],
            "projects": [],
            "certifications": [],
            "languages": []
        }

        if parsed:
            if parsed.experience:
                content["experience"] = parsed.experience
            if parsed.education:
                content["education"] = parsed.education
            if parsed.projects:
                content["projects"] = parsed.projects
            if parsed.certifications:
                content["certifications"] = parsed.certifications
            if parsed.languages:
                content["languages"] = parsed.languages

        return content

    def improve_text_with_ai(self, text: str, instruction: str) -> str:
        if not settings.GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY not configured")
        
        prompt = f"""
You are an expert Resume Writer.

Original Text:
"{text}"

Instruction / Goal:
{instruction}

Provide the rewritten or improved text directly, with NO conversational filler, markdown formatting, or quotes. Just the text.
"""
        response = self.model.generate_content(prompt)
        return response.text.strip()

    def get_drafts(self, db: Session, user_id: str):
        return db.query(ResumeDraft).filter(ResumeDraft.user_id == user_id).order_by(ResumeDraft.updated_at.desc()).all()

    def get_draft(self, db: Session, draft_id: str, user_id: str):
        return db.query(ResumeDraft).filter(ResumeDraft.id == draft_id, ResumeDraft.user_id == user_id).first()

    def save_draft(self, db: Session, user_id: str, draft_data: dict):
        draft_id = draft_data.get("id")
        
        if draft_id:
            draft = db.query(ResumeDraft).filter(ResumeDraft.id == draft_id, ResumeDraft.user_id == user_id).first()
            if draft:
                if "name" in draft_data: draft.name = draft_data["name"]
                if "template" in draft_data: draft.template = draft_data["template"]
                if "theme" in draft_data: draft.theme = draft_data["theme"]
                if "content" in draft_data: draft.content = draft_data["content"]
                
                db.commit()
                db.refresh(draft)
                return draft
                
        # Create new
        new_draft = ResumeDraft(
            user_id=user_id,
            name=draft_data.get("name", "Untitled Resume"),
            template=draft_data.get("template", "modern"),
            theme=draft_data.get("theme", "blue"),
            content=draft_data.get("content", {})
        )
        db.add(new_draft)
        db.commit()
        db.refresh(new_draft)
        return new_draft

    def generate_pdf(self, content: dict, template: str, theme: str):
        from xhtml2pdf import pisa
        from io import BytesIO
        
        name = content.get('personal_info', {}).get('name', 'Untitled')
        email = content.get('personal_info', {}).get('email', '')
        phone = content.get('personal_info', {}).get('phone', '')
        linkedin = content.get('personal_info', {}).get('linkedin', '')
        summary = content.get('summary', '')
        
        has_exp = content.get('hasExperience', True)
        experience = content.get('experience', []) if has_exp else []
        projects = content.get('projects', [])
        education = content.get('education', {})
        certifications = content.get('certifications', [])
        skills = content.get('skills', [])
        
        theme_color = "#2563eb" if theme == 'blue' else "#7c3aed" if theme == 'purple' else "#059669" if theme == 'green' else "#1f2937"
        
        html = f"""
        <html>
        <head>
            <style>
                @page {{ margin: 0.8in; }}
                body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; font-size: 11pt; color: #333; line-height: 1.4; }}
                h1 {{ font-size: 24pt; color: {theme_color}; margin-bottom: 5px; text-transform: uppercase; letter-spacing: 1px; text-align: center; }}
                .contact-info {{ text-align: center; font-size: 10pt; color: #555; margin-bottom: 20px; }}
                h2 {{ font-size: 13pt; color: {theme_color}; border-bottom: 1px solid #ccc; padding-bottom: 3px; margin-top: 15px; margin-bottom: 10px; text-transform: uppercase; }}
                .section-content {{ margin-bottom: 15px; }}
                .item-header {{ display: flex; justify-content: space-between; align-items: baseline; font-weight: bold; margin-bottom: 3px; }}
                .item-sub {{ display: flex; justify-content: space-between; font-size: 10pt; color: #555; margin-bottom: 5px; font-style: italic; }}
                .desc {{ font-size: 10pt; margin-top: 5px; white-space: pre-wrap; }}
                .skills-list {{ font-size: 10pt; }}
            </style>
        </head>
        <body>
            <h1>{name}</h1>
            <div class="contact-info">
                {email} | {phone} {f' | {linkedin}' if linkedin else ''}
            </div>
        """
        
        if summary:
            html += f"<h2>Professional Summary</h2><div class='section-content desc'>{summary}</div>"
            
        if projects:
            html += "<h2>Projects</h2><div class='section-content'>"
            for p in projects:
                html += f"""
                <div style="margin-bottom: 10px;">
                    <div class="item-header"><span>{p.get('title', '')} | <span style="font-weight:normal;">{p.get('technologies', '')}</span></span></div>
                    <div class="desc">{p.get('description', '')}</div>
                </div>
                """
            html += "</div>"
            
        if experience:
            html += "<h2>Experience</h2><div class='section-content'>"
            for ex in experience:
                html += f"""
                <div style="margin-bottom: 10px;">
                    <div class="item-header"><span>{ex.get('role', '')} at {ex.get('company', '')}</span><span>{ex.get('startDate', '')} - {'Present' if ex.get('current') else ex.get('endDate', '')}</span></div>
                    <div class="desc">{ex.get('description', '')}</div>
                </div>
                """
            html += "</div>"
            
        if education.get('degree') or education.get('twelfth') or education.get('tenth'):
            html += "<h2>Education</h2><div class='section-content'>"
            deg = education.get('degree', {})
            if deg.get('school'):
                html += f"""
                <div style="margin-bottom: 10px;">
                    <div class="item-header"><span>{deg.get('degree', '')} in {deg.get('field', '')}</span><span>{deg.get('startDate', '')} - {deg.get('endDate', '')}</span></div>
                    <div class="item-sub"><span>{deg.get('school', '')}</span><span>CGPA: {deg.get('cgpa', '')}</span></div>
                </div>
                """
            t12 = education.get('twelfth', {})
            if t12.get('school'):
                html += f"""
                <div style="margin-bottom: 10px;">
                    <div class="item-header"><span>Class XII ({t12.get('board', '')})</span><span>{t12.get('year', '')}</span></div>
                    <div class="item-sub"><span>{t12.get('school', '')}</span><span>{t12.get('percentage', '')}%</span></div>
                </div>
                """
            t10 = education.get('tenth', {})
            if t10.get('school'):
                html += f"""
                <div style="margin-bottom: 10px;">
                    <div class="item-header"><span>Class X ({t10.get('board', '')})</span><span>{t10.get('year', '')}</span></div>
                    <div class="item-sub"><span>{t10.get('school', '')}</span><span>{t10.get('percentage', '')}%</span></div>
                </div>
                """
            html += "</div>"
            
        if skills:
            html += f"<h2>Skills</h2><div class='section-content skills-list'>{', '.join(skills)}</div>"
            
        if certifications:
            html += "<h2>Certifications & Training</h2><div class='section-content'>"
            for c in certifications:
                html += f"""
                <div style="margin-bottom: 10px;">
                    <div class="item-header"><span>{c.get('name', '')} from {c.get('organization', '')}</span><span>{c.get('date', '')}</span></div>
                    <div class="desc">{c.get('description', '')}</div>
                </div>
                """
            html += "</div>"

        html += "</body></html>"
        
        pdf = BytesIO()
        pisa.CreatePDF(BytesIO(html.encode('utf-8')), dest=pdf)
        return pdf.getvalue()

    def generate_docx(self, content: dict, template: str = "modern"):
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        doc = Document()
        
        # Default styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        name = content.get('personal_info', {}).get('name', 'Untitled')
        email = content.get('personal_info', {}).get('email', '')
        phone = content.get('personal_info', {}).get('phone', '')
        
        # Header
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(name)
        run.bold = True
        run.font.size = Pt(24)
        
        contact = doc.add_paragraph(f"{email} | {phone}")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        def add_section_header(title):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(13)
            # Add a bottom border approximation
            p.paragraph_format.space_after = Pt(2)
        
        if content.get('summary'):
            add_section_header('Professional Summary')
            doc.add_paragraph(content.get('summary', ''))
            
        projects = content.get('projects', [])
        if projects:
            add_section_header('Projects')
            for p in projects:
                p_para = doc.add_paragraph()
                p_para.add_run(p.get('title', '')).bold = True
                p_para.add_run(f" | {p.get('technologies', '')}")
                doc.add_paragraph(p.get('description', ''))
                
        has_exp = content.get('hasExperience', True)
        experience = content.get('experience', []) if has_exp else []
        if experience:
            add_section_header('Experience')
            for ex in experience:
                p_para = doc.add_paragraph()
                p_para.add_run(f"{ex.get('role', '')} at {ex.get('company', '')}").bold = True
                dates = f"{ex.get('startDate', '')} - {'Present' if ex.get('current') else ex.get('endDate', '')}"
                p_para.add_run(f"  [{dates}]")
                doc.add_paragraph(ex.get('description', ''))
                
        education = content.get('education', {})
        if education.get('degree') or education.get('twelfth') or education.get('tenth'):
            add_section_header('Education')
            deg = education.get('degree', {})
            if deg.get('school'):
                doc.add_paragraph(f"{deg.get('degree', '')} in {deg.get('field', '')} ({deg.get('startDate', '')}-{deg.get('endDate', '')}) - {deg.get('school', '')} [CGPA: {deg.get('cgpa', '')}]")
            t12 = education.get('twelfth', {})
            if t12.get('school'):
                doc.add_paragraph(f"Class XII ({t12.get('board', '')}) - {t12.get('year', '')} - {t12.get('school', '')} [{t12.get('percentage', '')}%]")
            t10 = education.get('tenth', {})
            if t10.get('school'):
                doc.add_paragraph(f"Class X ({t10.get('board', '')}) - {t10.get('year', '')} - {t10.get('school', '')} [{t10.get('percentage', '')}%]")
                
        skills = content.get('skills', [])
        if skills:
            add_section_header('Skills')
            doc.add_paragraph(", ".join(skills))
            
        certifications = content.get('certifications', [])
        if certifications:
            add_section_header('Certifications & Training')
            for c in certifications:
                c_para = doc.add_paragraph()
                c_para.add_run(f"{c.get('name', '')} from {c.get('organization', '')}").bold = True
                doc.add_paragraph(c.get('description', ''))
        
        f = BytesIO()
        doc.save(f)
        return f.getvalue()

resume_builder_service = ResumeBuilderService()
