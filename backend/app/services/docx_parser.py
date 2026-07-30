import docx
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from paragraphs and tables in a DOCX file."""
    text_content = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def extract_structured_data_from_docx(file_path: str) -> list:
    """Extracts structured text blocks with paragraph IDs from a DOCX file."""
    pages = [] # DOCX doesn't have true pages, so we wrap it in one "page"
    blocks = []
    
    try:
        doc = docx.Document(file_path)
        
        for p_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                blocks.append({
                    "block_id": f"para_{p_idx}",
                    "text": text,
                    "bbox": [], # DOCX doesn't have bbox
                    "lines": [{"spans": [{"text": text}]}]
                })
                
        # Handle tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        blocks.append({
                            "block_id": f"table_{t_idx}_r{r_idx}_c{c_idx}",
                            "text": text,
                            "bbox": [],
                            "lines": [{"spans": [{"text": text}]}]
                        })
                        
        pages.append({
            "page_num": 0,
            "width": 800, # Mock width
            "height": 1000,
            "blocks": blocks
        })
        return pages
    except Exception as e:
        logger.error(f"Failed to parse structured DOCX {file_path}: {e}")
        raise ValueError(f"Failed to parse structured DOCX: {str(e)}")

import os
import tempfile

def apply_edits_to_docx(file_path: str, edits: list) -> str:
    """Applies text replacements to a DOCX and returns the new file path."""
    try:
        doc = docx.Document(file_path)
        for edit in edits:
            block_id = edit.get("block_id")
            new_text = edit.get("improved")
            
            if block_id and new_text:
                if block_id.startswith("para_"):
                    p_idx = int(block_id.split("_")[1])
                    if p_idx < len(doc.paragraphs):
                        # We just replace the text of the first run and clear others to maintain paragraph style
                        para = doc.paragraphs[p_idx]
                        if para.runs:
                            para.runs[0].text = new_text
                            for i in range(1, len(para.runs)):
                                para.runs[i].text = ""
                elif block_id.startswith("table_"):
                    parts = block_id.split("_")
                    t_idx = int(parts[1])
                    r_idx = int(parts[2][1:])
                    c_idx = int(parts[3][1:])
                    if t_idx < len(doc.tables):
                        table = doc.tables[t_idx]
                        if r_idx < len(table.rows):
                            row = table.rows[r_idx]
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.text = new_text
                                
        fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(temp_path)
        return temp_path
    except Exception as e:
        logger.error(f"Failed to apply edits to DOCX: {e}")
        raise ValueError(f"Failed to apply edits to DOCX: {str(e)}")
